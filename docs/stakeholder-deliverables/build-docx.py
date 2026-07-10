#!/usr/bin/env python3
"""Convert stakeholder Markdown sources in src/ to Word files in docx/."""

import re
import shutil
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "src"
DOCX = ROOT / "docx"
REFERENCE = ROOT / "reference.docx"

MAPPING = {
    "user-manual.md": "Bainum-User-Manual.docx",
    "faq.md": "Bainum-FAQ.docx",
    "terms-and-conditions.md": "Bainum-Terms-and-Conditions.docx",
    "usability-testing-protocol.md": "Bainum-Usability-Testing-Protocol.docx",
}


def _clean_inline_markdown(text: str) -> str:
    """Strip Markdown inline syntax so Word output is plain stakeholder text."""
    if not text:
        return text
    # [label](url) -> label
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # **bold** / __bold__
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    # `code`
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # stray markers (unbalanced ** from manual edits)
    return text.replace("**", "").replace("__", "").replace("`", "")


def build_with_pandoc(md_path: Path, out_path: Path) -> bool:
    cmd = ["pandoc", str(md_path), "-o", str(out_path)]
    if REFERENCE.exists():
        cmd.extend(["--reference-doc", str(REFERENCE)])
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(result.stderr, file=sys.stderr)
        return False
    return True


def _add_markdown_table(doc, header_cells, rows):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(header_cells):
        hdr[i].text = _clean_inline_markdown(text.strip())
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            if c_idx < len(cells):
                cells[c_idx].text = _clean_inline_markdown(text.strip())
    doc.add_paragraph()


def build_with_python_docx(md_path: Path, out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                rows = []
                for tl in table_lines[2:]:
                    if re.match(r"^[\|\s\-:]+$", tl):
                        continue
                    rows.append([c.strip() for c in tl.split("|")[1:-1]])
                if header_cells and rows:
                    _add_markdown_table(doc, header_cells, rows)
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = _clean_inline_markdown(stripped[level:].strip())
            doc.add_heading(title, level=min(level, 4))
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            for item in items:
                doc.add_paragraph(_clean_inline_markdown(item), style="List Bullet")
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            for item in items:
                doc.add_paragraph(_clean_inline_markdown(item), style="List Number")
            continue

        # Blockquote / emphasis line
        if stripped.startswith(">"):
            p = doc.add_paragraph(_clean_inline_markdown(stripped.lstrip("> ").strip()))
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        # Paragraph (collect until blank)
        para_lines = []
        while i < len(lines) and lines[i].strip() != "":
            if lines[i].strip().startswith("#"):
                break
            if lines[i].strip().startswith("|"):
                break
            if re.match(r"^[-*]\s+", lines[i].strip()):
                break
            if re.match(r"^\d+\.\s+", lines[i].strip()):
                break
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            doc.add_paragraph(_clean_inline_markdown(" ".join(para_lines)))

    doc.save(out_path)


def main() -> int:
    DOCX.mkdir(parents=True, exist_ok=True)
    use_pandoc = shutil.which("pandoc") is not None
    if use_pandoc:
        print("Using Pandoc for conversion.")
    else:
        print("Pandoc not found; using python-docx fallback.")
        try:
            import docx  # noqa: F401
        except ImportError:
            print("Install python-docx: pip3 install python-docx", file=sys.stderr)
            return 1

    for src_name, out_name in MAPPING.items():
        md_path = SRC / src_name
        if not md_path.exists():
            print(f"Missing source: {md_path}", file=sys.stderr)
            return 1
        out_path = DOCX / out_name
        ok = False
        if use_pandoc:
            ok = build_with_pandoc(md_path, out_path)
            if not ok:
                print(f"Pandoc failed for {src_name}; falling back to python-docx.")
        if not use_pandoc or not ok:
            build_with_python_docx(md_path, out_path)
        print(f"  {out_name}")

    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
