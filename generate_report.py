#!/usr/bin/env python3
"""
Generate annual report for Bainum Foundation
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call(["pip3", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('Bainum Project Annual Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# Introduction paragraph
intro = doc.add_paragraph()
intro.add_run('The Bainum Project Dashboard is a comprehensive web-based platform designed to support early childhood development assessment and tracking. The system enables teachers, administrators, and parents to collaboratively monitor children\'s language development through automated analysis of audio recordings. ').font.size = Pt(11)

# How it was built
heading1 = doc.add_heading('How It Was Built', level=1)
built_para = doc.add_paragraph()
built_para.add_run('The dashboard was developed using a modern full-stack architecture. The frontend is built with React and Vite, utilizing Tailwind CSS and DaisyUI for a responsive, accessible user interface. The backend is powered by Node.js and Express, with MongoDB serving as the database for storing child records, assessments, and user data. The system integrates with RevAI\'s Automatic Speech Recognition (ASR) API to transcribe audio recordings of children\'s speech. The application implements role-based access control with three distinct user types: administrators who manage the system, teachers who upload recordings and track student progress, and parents who can view their child\'s developmental data. Security features include JWT authentication, password hashing with bcrypt, and secure invitation-based registration for parents and teachers. ').font.size = Pt(11)

# What it does
heading2 = doc.add_heading('What It Does', level=1)
does_para1 = doc.add_paragraph()
does_para1.add_run('The dashboard provides a comprehensive suite of tools for tracking and analyzing children\'s language development. When teachers upload audio recordings, the system automatically transcribes the speech using RevAI and analyzes the transcripts for educational keywords across four key domains: Science Talk (scientific vocabulary and concepts), Social Talk (communication and interaction), Literature Talk (storytelling and narrative skills), and Language Development (overall language growth). ').font.size = Pt(11)

does_para2 = doc.add_paragraph()
does_para2.add_run('The platform visualizes this data through interactive dashboards featuring monthly dot matrix displays showing keyword frequency over time, as well as speedometer gauges that provide at-a-glance metrics for each developmental category. Teachers can review transcripts before accepting assessments, add observational notes, and track progress across multiple recordings. Parents receive secure, invitation-based access to view their child\'s developmental data, fostering transparency and engagement. Administrators can manage teacher and child profiles, view all transcripts, and access comprehensive analytics across the entire program. ').font.size = Pt(11)

# Future work
heading3 = doc.add_heading('Future Development', level=1)
future_para1 = doc.add_paragraph()
future_para1.add_run('Moving forward, we plan to enhance the dashboard with advanced AI-powered analysis capabilities. The current keyword-based approach will be augmented with semantic analysis using large language models to provide deeper insights into children\'s language development, including identification of developmental milestones, vocabulary complexity assessment, and personalized recommendations for teachers and parents. ').font.size = Pt(11)

future_para2 = doc.add_paragraph()
future_para2.add_run('Additional planned features include longitudinal trend analysis with predictive modeling, enhanced data visualization with comparative analytics across children and classrooms, automated report generation for stakeholders, and mobile application support for easier recording uploads in classroom settings. We also aim to integrate with existing educational assessment frameworks and expand the keyword taxonomy based on research findings and user feedback. ').font.size = Pt(11)

# Save document
output_path = '/home/vashist/Desktop/Desktop/Anita Zucker Center/Bainum Project/Main/Bainum_Project_Annual_Report.docx'
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
